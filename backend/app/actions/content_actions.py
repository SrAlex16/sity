from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from elevenlabs import save as elevenlabs_save

from app.memory.db import get_session
from app.memory.models import Episode, NewsItem

_PROJECT_ROOT = Path(__file__).resolve().parents[3]
_PROMPTS_DIR = _PROJECT_ROOT / "config" / "prompts"

_LARGO_PROMPT_PATH = _PROMPTS_DIR / "script_largo_prompt.txt"
_SHORTS_PROMPT_PATH = _PROMPTS_DIR / "script_shorts_prompt.txt"

# Structural labels that should not be narrated
_TTS_SKIP_PREFIXES = (
    "SITY:", "**SITY:**", "*SITY:*", "SITY**:",
    "[NOTA PRODUCCIÓN", "*[NOTA",
    "GUION", "Generado:", "---",
    "**Duración estimada", "Duración estimada",
)
_TTS_SKIP_PATTERNS = (
    "GANCHO (", "DESARROLLO:", "CIERRE:",
    "INTRO (", "SECCIÓN ", "SHORT ", "OUTRO (", "REFLEXIÓN FINAL",
)


def _clean_for_tts(text: str) -> str:
    text = re.sub(r'\*\*.*?\*\*', '', text)
    text = re.sub(r'\*.*?\*', '', text)
    text = text.replace('*', '').replace('---', '')
    text = re.sub(r' +', ' ', text).strip()
    return text


def _expand_acronyms_with_claude(text: str, api_key: str) -> str:
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=len(text) + 500,
        messages=[{
            "role": "user",
            "content": (
                "El siguiente texto va a ser narrado por una IA de texto a voz "
                "en español de España. Reescribe el texto expandiendo los acrónimos "
                "y siglas para que suenen naturales al escucharlos (por ejemplo, "
                "'HBO' → 'Hachebeo', 'IA' → 'Inteligencia Artificial' si aparece "
                "por primera vez, etc.). No cambies nada más del texto — solo los "
                "acrónimos y siglas que sonarían raros al leerlos en voz alta. "
                "Devuelve únicamente el texto corregido, sin explicaciones.\n\n"
                f"{text}"
            ),
        }],
    )
    block = response.content[0]
    return block.text if hasattr(block, "text") else text  # type: ignore[union-attr]


@dataclass
class ContentActionResult:
    ok: bool
    text: str


def execute_content_action(payload: dict[str, Any]) -> ContentActionResult:
    action = payload.get("action")
    if action == "select_news":
        return _execute_select_news(payload)
    if action == "generate_script":
        return _execute_generate_script(payload)
    if action == "generate_tts":
        return _execute_generate_tts(payload)
    if action == "generate_images":
        return _execute_generate_images(payload)
    return ContentActionResult(ok=False, text=f"Acción desconocida: {action}")


def parse_payload(payload_json: str) -> dict[str, Any]:
    return json.loads(payload_json)


def _execute_select_news(payload: dict[str, Any]) -> ContentActionResult:
    from sqlmodel import select as sql_select

    session = next(get_session())
    news_ids = payload.get("news_ids", [])
    status = payload.get("status", "selected")
    updated = 0
    for nid in news_ids:
        item = session.exec(sql_select(NewsItem).where(NewsItem.id == nid)).first()
        if item:
            item.status = status
            updated += 1
    session.commit()
    return ContentActionResult(ok=True, text=f"{updated} noticia(s) marcadas como '{status}'.")


def _build_docx(text: str, output_path: Path, title: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    heading = doc.add_heading(f"GUION — Sity Canal · {title}", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        else:
            p = doc.add_paragraph(stripped)
            if "[NOTA PRODUCCIÓN:" in stripped or stripped.startswith("*[NOTA"):
                for run in p.runs:
                    run.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _execute_generate_script(payload: dict[str, Any]) -> ContentActionResult:
    import anthropic
    from sqlmodel import select as sql_select

    news_ids = payload.get("news_ids", [])
    output_dir = Path(payload.get("output_dir", "work/canal/guiones"))

    session = next(get_session())
    date_str = datetime.now().strftime("%Y-%m-%d")
    episode = Episode(status="draft")
    session.add(episode)
    session.flush()
    ep_id = episode.id
    ep_label = f"EP{ep_id:03d}"

    largo_path = output_dir / f"{ep_label}-largo-{date_str}.docx"
    shorts_path = output_dir / f"{ep_label}-shorts-{date_str}.docx"

    # Build shared news text
    news_items_text = payload.get("news_items_text", "")

    largo_template = _LARGO_PROMPT_PATH.read_text(encoding="utf-8")
    shorts_template = _SHORTS_PROMPT_PATH.read_text(encoding="utf-8")
    largo_prompt = largo_template.replace("{news_items}", news_items_text)
    shorts_prompt = shorts_template.replace("{news_items}", news_items_text)

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    largo_response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        messages=[{"role": "user", "content": largo_prompt}],
    )
    largo_block = largo_response.content[0]
    largo_text: str = largo_block.text if hasattr(largo_block, "text") else ""  # type: ignore[union-attr]

    shorts_response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        messages=[{"role": "user", "content": shorts_prompt}],
    )
    shorts_block = shorts_response.content[0]
    shorts_text: str = shorts_block.text if hasattr(shorts_block, "text") else ""  # type: ignore[union-attr]

    _build_docx(largo_text, largo_path, f"{ep_label} — Largo")
    _build_docx(shorts_text, shorts_path, f"{ep_label} — Shorts")

    episode.script_path = str(largo_path)
    episode.script_shorts_path = str(shorts_path)
    episode.status = "script_ready"

    for nid in news_ids:
        item = session.exec(sql_select(NewsItem).where(NewsItem.id == nid)).first()
        if item:
            item.status = "used"
            item.episode_id = ep_id
    session.commit()

    return ContentActionResult(
        ok=True,
        text=(
            f"Guion generado: {ep_label}\n"
            f"Largo: {largo_path}\n"
            f"Shorts: {shorts_path}\n"
            f"{len(news_ids)} noticia(s) vinculadas al episodio.\n"
            f"Abre los archivos para revisarlos antes de continuar."
        ),
    )


def _execute_generate_tts(payload: dict[str, Any]) -> ContentActionResult:
    from elevenlabs.client import ElevenLabs
    from elevenlabs.types import VoiceSettings
    from docx import Document
    from sqlmodel import select as sql_select

    episode_id = payload.get("episode_id")
    script_path = Path(payload.get("script_path", ""))
    audio_path = Path(payload.get("audio_path", ""))
    script_type = str(payload.get("script_type", "largo"))

    if not script_path.exists():
        return ContentActionResult(ok=False, text=f"No se encontró el guion en {script_path}")

    doc = Document(str(script_path))
    narrable_lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if any(text.startswith(p) for p in _TTS_SKIP_PREFIXES):
            continue
        if any(p in text for p in _TTS_SKIP_PATTERNS):
            continue
        if para.style and para.style.name.startswith("Heading"):
            continue
        clean = _clean_for_tts(text)
        if clean:
            narrable_lines.append(clean)

    full_text = "\n".join(narrable_lines)
    if not full_text.strip():
        return ContentActionResult(ok=False, text="El guion no tiene texto narrable extraíble.")

    anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")
    if anthropic_key:
        try:
            full_text = _expand_acronyms_with_claude(full_text, anthropic_key)
        except Exception:
            pass  # fall back to original text

    api_key = os.getenv("ELEVENLABS_API_KEY", "")
    voice_id = os.getenv("ELEVENLABS_VOICE_ID", "")
    if not api_key or not voice_id:
        return ContentActionResult(ok=False, text="Faltan ELEVENLABS_API_KEY o ELEVENLABS_VOICE_ID en .env")

    try:
        client = ElevenLabs(api_key=api_key)
        audio = client.text_to_speech.convert(
            voice_id=voice_id,
            text=full_text,
            model_id="eleven_multilingual_v2",
            voice_settings=VoiceSettings(stability=0.5, similarity_boost=0.75),
        )
        audio_path.parent.mkdir(parents=True, exist_ok=True)
        elevenlabs_save(audio, str(audio_path))
    except Exception as e:
        return ContentActionResult(ok=False, text=f"Error llamando a ElevenLabs: {e}")

    session = next(get_session())
    episode = session.exec(sql_select(Episode).where(Episode.id == episode_id)).first()
    if episode:
        if script_type == "shorts":
            episode.audio_shorts_path = str(audio_path)
        else:
            episode.audio_path = str(audio_path)
            episode.status = "audio_ready"
        session.add(episode)
        session.commit()

    size_mb = audio_path.stat().st_size / (1024 * 1024)
    type_label = "shorts" if script_type == "shorts" else "largo"
    return ContentActionResult(
        ok=True,
        text=(
            f"Audio {type_label} generado: {audio_path}\n"
            f"Tamaño: {size_mb:.1f} MB\n"
            f"Episodio actualizado (audio_{type_label}_path).\n"
            f"Escucha el archivo antes de continuar."
        ),
    )


def _execute_generate_images(payload: dict[str, Any]) -> ContentActionResult:
    import anthropic
    import httpx
    from sqlmodel import Session, select as sql_select

    from app.memory.db import engine
    from app.memory.models import Episode

    episode_id = payload.get("episode_id")
    ep_label = f"EP{episode_id:03d}" if episode_id else Path(payload.get("assets_dir", "EP000")).name
    timestamps_dir = _PROJECT_ROOT / "work" / "canal" / "guiones" / "timestamps"
    transcript_path = timestamps_dir / f"{ep_label}.txt"
    assets_dir = _PROJECT_ROOT / "work" / "canal" / "assets" / ep_label

    if not transcript_path.exists():
        return ContentActionResult(
            ok=False, text=f"No se encontró la transcripción en {transcript_path}"
        )

    raw_text = transcript_path.read_text(encoding="utf-8")

    raw_text = re.sub(
        r'\(Transcrito por TurboScribe.*?\)',
        '',
        raw_text,
        flags=re.IGNORECASE,
    ).strip()

    parts = re.compile(r'\((\d+:\d+)\)\s*').split(raw_text)
    # parts: ['', '0:06', 'texto1', '0:12', 'texto2', ...]
    segments: list[tuple[str, str]] = []
    i = 1
    while i < len(parts) - 1:
        timestamp = parts[i].strip()
        text = parts[i + 1].strip()
        if text:
            segments.append((timestamp, text))
        i += 2

    if not segments:
        return ContentActionResult(
            ok=False,
            text=(
                "No se encontraron segmentos con timestamps válidos. "
                "Formato esperado: (0:00) texto (0:06) más texto..."
            ),
        )

    assets_dir.mkdir(parents=True, exist_ok=True)

    anthropic_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY", ""))
    stability_api_key = os.getenv("STABILITY_API_KEY", "")

    if not stability_api_key:
        return ContentActionResult(ok=False, text="Falta STABILITY_API_KEY en el entorno.")

    results: list[str] = []
    errors: list[str] = []

    for i, (timestamp, text) in enumerate(segments, start=1):
        text = text.strip()
        img_filename = f"{ep_label}-img-{i:03d}.png"
        img_path = assets_dir / img_filename

        try:
            prompt_response = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=300,
                messages=[{
                    "role": "user",
                    "content": (
                        "Genera un prompt en inglés para Stable Diffusion "
                        "que ilustre visualmente el siguiente fragmento de "
                        "narración de un canal de divulgación tech/IA.\n\n"
                        f"Texto narrado: \"{text}\"\n\n"
                        "El prompt debe seguir esta estética SIEMPRE:\n"
                        "- Estilo cyberpunk con influencia asiática\n"
                        "- Predominio de colores: azul eléctrico, magenta, "
                        "rosa neón, violeta, cian\n"
                        "- Interfaces holográficas y pantallas flotantes\n"
                        "- Fondos oscuros con neones brillantes y contrastes fuertes\n"
                        "- Inspirado en Blade Runner, Ghost in the Shell, Cyberpunk 2077\n"
                        "- Composición simple y limpia, no sobrecargada\n"
                        "- Sin texto en la imagen excepto caracteres "
                        "decorativos asiáticos si encajan\n"
                        "- Formato 16:9 horizontal\n\n"
                        "Devuelve SOLO el prompt en inglés, sin explicaciones."
                    ),
                }],
            )
            prompt_block = prompt_response.content[0]
            image_prompt: str = prompt_block.text.strip() if hasattr(prompt_block, "text") else ""  # type: ignore[union-attr]
        except Exception as e:
            errors.append(f"[{timestamp}] Error generando prompt: {e}")
            continue

        try:
            response = httpx.post(
                "https://api.stability.ai/v2beta/stable-image/generate/sd3",
                headers={
                    "authorization": f"Bearer {stability_api_key}",
                    "accept": "image/*",
                },
                data={
                    "prompt": image_prompt,
                    "aspect_ratio": "16:9",
                    "model": "sd3.5-medium",
                    "output_format": "png",
                },
                timeout=60,
            )
            if response.status_code == 200:
                img_path.write_bytes(response.content)
                results.append(f"✓ ({timestamp}) → {img_filename}")
            else:
                errors.append(
                    f"[{timestamp}] Error Stability AI: "
                    f"{response.status_code} {response.text[:200]}"
                )
        except Exception as e:
            errors.append(f"[{timestamp}] Error llamando a Stability AI: {e}")

    if episode_id and len(results) > 0:
        with Session(engine) as session:
            episode = session.exec(
                sql_select(Episode).where(Episode.id == episode_id)
            ).first()
            if episode:
                session.add(episode)
                session.commit()

    summary = (
        f"Generación completada para {ep_label}:\n"
        f"✓ {len(results)} imágenes generadas\n"
        f"✗ {len(errors)} errores\n\n"
    )
    if results:
        summary += "Imágenes generadas:\n" + "\n".join(results)
    if errors:
        summary += "\n\nErrores:\n" + "\n".join(errors)
    summary += f"\n\nImágenes guardadas en: {assets_dir}"

    return ContentActionResult(ok=len(results) > 0, text=summary)
