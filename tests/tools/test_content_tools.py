"""Tests para las tools del canal de YouTube.

Todos los tests usan mocks — sin llamadas reales a RSS ni a Anthropic.
"""
from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from app.tools.registry import ToolContext


def _make_ctx(tool_name: str, tool_input: dict) -> ToolContext:
    executor = MagicMock()
    executor.session = MagicMock()
    return ToolContext(
        tool_name=tool_name,
        tool_input=tool_input,
        trace_id="test-content",
        executor=executor,
    )


def _make_news_item(id: int, title: str, source: str, status: str = "pending") -> MagicMock:
    item = MagicMock()
    item.id = id
    item.title = title
    item.source = source
    item.category = "tech"
    item.summary = "Resumen de prueba"
    item.url = f"https://example.com/news/{id}"
    item.published_at = "2026-07-01T10:00:00+00:00"
    item.status = status
    item.created_at = "2026-07-01T10:00:00"
    return item


# ── fetch_rss_news ────────────────────────────────────────────────────────────

def test_fetch_rss_news_saves_to_db() -> None:
    from app.tools.handlers.content_tools import handle_fetch_rss_news

    mock_feed = MagicMock()
    mock_entry = MagicMock()
    mock_entry.get = lambda k, d="": {"title": "Test Entry", "link": "https://example.com/1", "summary": "Un resumen"}.get(k, d)
    mock_entry.published_parsed = (2026, 7, 1, 10, 0, 0, 0, 0, 0)
    mock_feed.entries = [mock_entry]

    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = None  # no duplicate
    mock_session.exec.return_value.all.return_value = []

    ctx = _make_ctx("fetch_rss_news", {})

    canal_cfg = {
        "rss_feeds": [{"name": "Test Feed", "url": "https://example.com/feed", "category": "tech"}],
        "settings": {"days_back": 7, "output": {}},
    }

    with patch("app.tools.handlers.content_tools._load_canal_config", return_value=canal_cfg):
        with patch("feedparser.parse", return_value=mock_feed):
            with patch("app.tools.handlers.content_tools.get_session", return_value=iter([mock_session])):
                result = handle_fetch_rss_news(ctx)

    assert result.ok is True
    assert "Ingesta completada" in result.message
    mock_session.add.assert_called_once()
    mock_session.commit.assert_called()


def test_fetch_rss_news_deduplicates_urls() -> None:
    from app.tools.handlers.content_tools import handle_fetch_rss_news

    mock_feed = MagicMock()
    mock_entry = MagicMock()
    mock_entry.get = lambda k, d="": {"title": "Dup", "link": "https://example.com/dup", "summary": ""}.get(k, d)
    mock_entry.published_parsed = (2026, 7, 1, 10, 0, 0, 0, 0, 0)
    mock_feed.entries = [mock_entry]

    existing_item = MagicMock()
    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = existing_item  # duplicate found
    mock_session.exec.return_value.all.return_value = []

    ctx = _make_ctx("fetch_rss_news", {})
    canal_cfg = {
        "rss_feeds": [{"name": "Feed", "url": "https://x.com/feed", "category": "tech"}],
        "settings": {"days_back": 7, "output": {}},
    }

    with patch("app.tools.handlers.content_tools._load_canal_config", return_value=canal_cfg):
        with patch("feedparser.parse", return_value=mock_feed):
            with patch("app.tools.handlers.content_tools.get_session", return_value=iter([mock_session])):
                result = handle_fetch_rss_news(ctx)

    assert result.ok is True
    assert "duplicada" in result.message
    mock_session.add.assert_not_called()


# ── select_news ───────────────────────────────────────────────────────────────

def test_select_news_creates_pending_action() -> None:
    from app.tools.handlers.content_tools import handle_select_news

    ctx = _make_ctx("select_news", {"news_ids": [1, 3, 5], "action": "selected"})

    mock_manager = MagicMock()
    mock_manager.find_equivalent_pending_action.return_value = None
    created = MagicMock()
    created.id = "act_content01"
    created.confirmation_phrase = "confirmo ejecutar act_content01"
    created.summary = "Marcar 3 noticia(s) como 'selected': IDs [1, 3, 5]"
    mock_manager.create_pending_action.return_value = created

    with patch("app.tools.handlers.content_tools.ConfirmationManager", return_value=mock_manager):
        result = handle_select_news(ctx)

    assert result.ok is True
    assert "confirmo ejecutar act_content01" in result.message
    assert result.raw_result.get("local_final") is True
    mock_manager.create_pending_action.assert_called_once()


def test_select_news_missing_ids_returns_error() -> None:
    from app.tools.handlers.content_tools import handle_select_news

    ctx = _make_ctx("select_news", {"action": "selected"})
    result = handle_select_news(ctx)

    assert result.ok is False
    assert "news_ids" in result.message


def test_select_news_invalid_action_returns_error() -> None:
    from app.tools.handlers.content_tools import handle_select_news

    ctx = _make_ctx("select_news", {"news_ids": [1], "action": "invalid"})
    result = handle_select_news(ctx)

    assert result.ok is False
    assert "selected" in result.message or "discarded" in result.message


# ── generate_script ───────────────────────────────────────────────────────────

def test_generate_script_fails_without_selected_news() -> None:
    from app.tools.handlers.content_tools import handle_generate_script

    mock_session = MagicMock()
    mock_session.exec.return_value.all.return_value = []

    ctx = _make_ctx("generate_script", {})

    with patch("app.tools.handlers.content_tools.get_session", return_value=iter([mock_session])):
        result = handle_generate_script(ctx)

    assert result.ok is False
    assert "seleccionadas" in result.message.lower()


def test_generate_script_creates_pending_action() -> None:
    from app.tools.handlers.content_tools import handle_generate_script

    selected = [_make_news_item(1, "Noticia A", "The Verge", "selected")]
    mock_session = MagicMock()
    mock_session.exec.return_value.all.return_value = selected

    ctx = _make_ctx("generate_script", {})

    mock_manager = MagicMock()
    mock_manager.find_equivalent_pending_action.return_value = None
    created = MagicMock()
    created.id = "act_gen01"
    created.confirmation_phrase = "confirmo ejecutar act_gen01"
    created.summary = "Generar guion con 1 noticia(s)"
    mock_manager.create_pending_action.return_value = created

    canal_cfg = {"settings": {"output": {"guiones": "work/canal/guiones"}}}

    with patch("app.tools.handlers.content_tools.get_session", return_value=iter([mock_session])):
        with patch("app.tools.handlers.content_tools.ConfirmationManager", return_value=mock_manager):
            with patch("app.tools.handlers.content_tools._load_canal_config", return_value=canal_cfg):
                with patch("pathlib.Path.read_text", return_value="Prompt: {news_items}"):
                    result = handle_generate_script(ctx)

    assert result.ok is True
    assert "confirmo ejecutar act_gen01" in result.message
    assert result.raw_result.get("local_final") is True


# ── content_actions ───────────────────────────────────────────────────────────

def test_execute_select_news_updates_status() -> None:
    from app.actions.content_actions import _execute_select_news

    item = MagicMock()
    item.status = "pending"

    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = item

    with patch("app.actions.content_actions.get_session", return_value=iter([mock_session])):
        result = _execute_select_news({"news_ids": [42], "status": "selected"})

    assert result.ok is True
    assert item.status == "selected"
    assert "1" in result.text
    mock_session.commit.assert_called_once()


def test_execute_generate_script_creates_docx(tmp_path: Path) -> None:
    from app.actions.content_actions import _execute_generate_script
    from app.memory.models import Episode

    mock_response = MagicMock()
    mock_response.content = [MagicMock(text="# Título\n\nContenido del guion.")]

    mock_client = MagicMock()
    mock_client.messages.create.return_value = mock_response

    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = MagicMock()

    def mock_flush_sets_id() -> None:
        # Simulate DB assigning an id on flush by setting id on any Episode added
        for call in mock_session.add.call_args_list:
            obj = call[0][0]
            if isinstance(obj, Episode):
                obj.id = 1

    mock_session.flush.side_effect = mock_flush_sets_id

    with patch("anthropic.Anthropic", return_value=mock_client):
        with patch("app.actions.content_actions.get_session", return_value=iter([mock_session])):
            result = _execute_generate_script({
                "full_prompt": "Genera un guion sobre: {news_items}",
                "output_dir": str(tmp_path),
                "news_ids": [1],
            })

    assert result.ok is True
    expected_docx = list(tmp_path.glob("EP001-*.docx"))
    assert expected_docx, f"No EP001-*.docx found in {tmp_path}"
    assert "EP001" in result.text


# ── generate_tts ──────────────────────────────────────────────────────────────

def test_generate_tts_no_script_ready_fails() -> None:
    from app.tools.handlers.content_tools import handle_generate_tts

    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = None

    ctx = _make_ctx("generate_tts", {})

    with patch("app.tools.handlers.content_tools.get_session", return_value=iter([mock_session])):
        result = handle_generate_tts(ctx)

    assert result.ok is False
    assert "script_ready" in result.message


def test_generate_tts_creates_pending_action() -> None:
    from app.tools.handlers.content_tools import handle_generate_tts
    from app.memory.models import Episode

    episode = MagicMock(spec=Episode)
    episode.id = 1
    episode.status = "script_ready"
    episode.script_path = "/tmp/EP001-2026-07-02.docx"

    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = episode

    ctx = _make_ctx("generate_tts", {})

    mock_manager = MagicMock()
    mock_manager.find_equivalent_pending_action.return_value = None
    created = MagicMock()
    created.id = "act_tts01"
    created.confirmation_phrase = "confirmo ejecutar act_tts01"
    created.summary = "Generar audio TTS de EP001 → EP001.mp3"
    mock_manager.create_pending_action.return_value = created

    canal_cfg = {"settings": {"output": {"audio": "work/canal/audio"}}}

    with patch("app.tools.handlers.content_tools.get_session", return_value=iter([mock_session])):
        with patch("app.tools.handlers.content_tools.ConfirmationManager", return_value=mock_manager):
            with patch("app.tools.handlers.content_tools._load_canal_config", return_value=canal_cfg):
                result = handle_generate_tts(ctx)

    assert result.ok is True
    assert "confirmo ejecutar act_tts01" in result.message
    assert result.raw_result.get("local_final") is True


def test_execute_generate_tts_extracts_narrable_text(tmp_path: Path) -> None:
    """El extractor debe omitir notas de producción, encabezados y metadatos."""
    from docx import Document
    from app.actions.content_actions import _execute_generate_tts

    script_path = tmp_path / "EP001-test.docx"
    audio_path = tmp_path / "EP001.mp3"

    doc = Document()
    doc.add_heading("GUION — Sity Canal · EP001", 0)
    doc.add_paragraph("Generado: 2026-07-02 15:00")
    doc.add_heading("Sección principal", level=1)
    doc.add_paragraph("Este es el texto narrable del episodio.")
    doc.add_paragraph("[NOTA PRODUCCIÓN: pausa aquí, mostrar b-roll]")
    doc.add_paragraph("Otro párrafo narrable sin notas.")
    doc.save(str(script_path))

    mock_audio = iter([b"fake_audio_bytes"])
    mock_client = MagicMock()
    mock_client.text_to_speech.convert.return_value = mock_audio

    mock_session = MagicMock()
    mock_session.exec.return_value.first.return_value = MagicMock()

    def mock_save(audio, filename: str) -> None:
        Path(filename).write_bytes(b"fake_audio")

    with patch("elevenlabs.client.ElevenLabs", return_value=mock_client):
        with patch("app.actions.content_actions.elevenlabs_save", mock_save):
            with patch("app.actions.content_actions.get_session", return_value=iter([mock_session])):
                with patch.dict("os.environ", {"ELEVENLABS_API_KEY": "test", "ELEVENLABS_VOICE_ID": "v1"}):
                    result = _execute_generate_tts({
                        "episode_id": 1,
                        "script_path": str(script_path),
                        "audio_path": str(audio_path),
                    })

    assert result.ok is True
    call_kwargs = mock_client.text_to_speech.convert.call_args
    sent_text: str = call_kwargs.kwargs.get("text") or call_kwargs.args[1] if call_kwargs.args else ""
    if not sent_text:
        sent_text = call_kwargs.kwargs["text"]
    assert "[NOTA PRODUCCIÓN:" not in sent_text
    assert "GUION" not in sent_text
    assert "Este es el texto narrable" in sent_text
    assert "Otro párrafo narrable" in sent_text
